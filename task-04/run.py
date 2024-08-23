import json
import docx

from GoogleBooks import GoogleBooksAPI

from telebot import TeleBot
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton

fp = open('config.json', 'r')
config = json.load(fp)

bot = TeleBot(config['telegram_config']['bot_key'])
gb_api = GoogleBooksAPI()

@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.send_message(
        message.chat.id, "Hey, I'm PagePal! A Book recommendation telegram bot", parse_mode='Markdown'
    )

@bot.message_handler(commands=['help'])
def send_help(message):
    msg = ""
    msg += "1. /start returns a welcome message to the user\n"
    msg += "2. /book the user is asked to enter the genre name, upon which the user types in the name and the bot returns a CSV file with different book names and their details.\n"
    msg += "3. /preview  the user is asked to enter the book name, upon which it redirects us to the preview link.\n"
    msg += "4. /list asks the user to type in a specific book name and then it returns a message to execute /reading_list command\n"
    msg += "5. /reading_list displays three buttons\n"
    msg += "    a. Add a book - Press this button to add a new book to the reading list\n"
    msg += "    b. Delete a book - Press this button to remove a book from the reading list\n"
    msg += "    c. View Reading List - Press this button to return the reading list (docx)\n"
    msg += "6. /help returns the list of commands with their description"

    bot.send_message(message.chat.id, msg, parse_mode="Markdown")

@bot.message_handler(commands=['book'])
def send_books(message):
    msg = bot.send_message(
        message.chat.id, "Please enter genre: "
    )

    bot.register_next_step_handler(msg, process_genre)

@bot.message_handler(commands=['preview'])
def send_preview(message):
    msg = bot.send_message(
        message.chat.id, "Please enter book name: "
    )

    bot.register_next_step_handler(msg, process_preview)

@bot.message_handler(commands=['list'])
def send_list(message):
    msg = bot.send_message(
        message.chat.id, "Please enter your book name: ", parse_mode="Markdown"
    )

    bot.register_next_step_handler(msg, process_book_name)

@bot.message_handler(commands=['reading_list'])
def send_reading_list(message):
    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("Add a book", callback_data='add'))
    markup.add(InlineKeyboardButton("Delete a book", callback_data='delete'))
    markup.add(InlineKeyboardButton("View Reading List", callback_data='view'))

    bot.send_message(
        message.chat.id, "You can manage your reading list here:", reply_markup=markup
    )

@bot.callback_query_handler(func=lambda btn: True)
def handle_button(btn):
    if btn.data == 'add':
        book_name = bot.user_data.get('book_name', '')

        if book_name:
            gb_api.reading_list.append(book_name)
            bot.send_message(btn.message.chat.id, f'"{book_name}" is added to your reading list.')
        else:
            bot.send_message(btn.message.chat.id, 'No book name is found. Use /list to enter a book name first.')

    elif btn.data == 'delete':
        book_name = bot.user_data.get('book_name', '')

        if book_name in gb_api.reading_list:
            gb_api.reading_list.remove(book_name)
            bot.send_message(btn.message.chat.id, f'"{book_name}" is removed from your reading list.')
        else:
            bot.send_message(btn.message.chat.id, f'"{book_name}" is not found in your reading list.')

    elif btn.data == 'view':
        doc = docx.Document()
        doc.add_heading('Reading List', 0)

        for book in gb_api.reading_list:
            doc.add_paragraph(book)

        file_name = 'reading_list.docx'
        doc.save(file_name)

        fp = open(file_name, 'rb')
        bot.send_document(btn.message.chat.id, fp, caption="Download your reading list:")

def process_genre(message):
    genre_name = message.text

    books, code = gb_api.get_books_genre(genre_name)
    if code == 200:
        path = gb_api.convert_books_to_csv(books, genre_name)
        fp = open(path, 'r')
        bot.send_document(message.chat.id, fp, caption=f"Download {genre_name} books list:")

    else:
        bot.send_message(message.chat.id, f"Unable to get books, ERR_CODE: {code}")

def process_preview(message):
    book_name = message.text

    link, code = gb_api.get_preview_link(book_name)
    if code == 200:
        bot.send_message(
            message.chat.id, f"Link for {book_name}\n{link}"
        )

    else:
        bot.send_message(
            message.chat.id, f"Unable to get links, ERR_CODE: {code}"
        )

def process_book_name(message):
    book_name = message.text

    bot.send_message(
        message.chat.id, f"{book_name} is added to reading list.\nPlease enter /reading_list to manage your reading list."
    )
    bot.user_data = {'book_name': book_name}

    gb_api.reading_list.append(book_name)

if __name__ == '__main__':
    bot.infinity_polling()